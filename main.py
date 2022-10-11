import wikipedia
from docx import Document
from docx.shared import Pt
from os import system

wikipedia.set_lang('PT')

def do_work(name, theme):
    document = Document()
    document.add_heading(name, 2,)

    title = document.add_heading(theme, 0)
    title.alignment = 1

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    r = get_info(theme)
    p = document.add_paragraph(r)
    p.style = document.styles['Normal']

    #document.add_picture('name.png', width=Inches(1.25))
    
    document.save('homework.docx')

def get_info(theme):
    # request
    r = wikipedia.summary(theme)
    return r

def gui():
    while True:
        system('cls')
        print('==== Criar homework ====')
        name = input('Seu nome: ')
        theme = input('Tema: ')

        print('Criando homework...')
        do_work(name, theme)
        print('==== Sucesso ao criar. ====')
        system('pause')

gui()